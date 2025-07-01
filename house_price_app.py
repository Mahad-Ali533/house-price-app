import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.model_selection import train_test_split
from sklearn.linear_model import LinearRegression
from sklearn.metrics import mean_squared_error, r2_score
from sklearn.preprocessing import StandardScaler
from sklearn.datasets import fetch_california_housing
from io import BytesIO
from docx import Document

# === Data Generation & Loading ===

def generate_dataset(n_samples, min_sqft, max_sqft):
    np.random.seed(42)
    bedrooms = np.random.randint(1, 6, size=n_samples)
    sqft = np.random.randint(min_sqft, max_sqft, size=n_samples)
    noise = np.random.normal(0, 20000, size=n_samples)
    prices = 50000 + 30000 * bedrooms + 100 * sqft + noise
    return pd.DataFrame({'Bedrooms': bedrooms, 'SquareFootage': sqft, 'Price': prices.astype(int)})

def load_real_dataset():
    housing = fetch_california_housing(as_frame=True)
    df = housing.frame
    df = df.rename(columns={'AveRooms': 'Bedrooms', 'MedHouseVal': 'Price'})
    df = df[['Bedrooms', 'HouseAge', 'Price']]
    df['Price'] *= 100000
    return df

# === Model Training ===

def train_model(df):
    feature_cols = [col for col in df.columns if col != 'Price']
    X = df[feature_cols]
    y = df['Price']
    scaler = StandardScaler()
    X_scaled = scaler.fit_transform(X)
    X_train, X_test, y_train, y_test = train_test_split(X_scaled, y, test_size=0.2, random_state=42)
    model = LinearRegression()
    model.fit(X_train, y_train)
    y_pred = model.predict(X_test)
    mse = mean_squared_error(y_test, y_pred)
    r2 = r2_score(y_test, y_pred)
    return model, X_test, y_test, y_pred, mse, r2

# === Plot ===

def plot_results(y_test, y_pred):
    fig, ax = plt.subplots()
    sns.scatterplot(x=y_test, y=y_pred, ax=ax)
    ax.plot([y_test.min(), y_test.max()], [y_test.min(), y_test.max()], 'r--')
    ax.set_xlabel("Actual Price")
    ax.set_ylabel("Predicted Price")
    ax.set_title("Actual vs Predicted House Prices")
    return fig

# === Word Report ===

def generate_docx_report(mse, r2):
    doc = Document()
    doc.add_heading("House Price Prediction Report", 0)
    doc.add_heading("Evaluation Metrics", level=1)
    doc.add_paragraph(f"Mean Squared Error: {mse:.2f}")
    doc.add_paragraph(f"R² Score: {r2:.4f}")
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# === Excel Export ===

def generate_excel(y_test, y_pred, mse, r2):
    output = BytesIO()
    df = pd.DataFrame({'Actual Price': y_test.values, 'Predicted Price': y_pred})
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Predictions')
        summary = pd.DataFrame({'Metric': ['MSE', 'R2'], 'Value': [mse, r2]})
        summary.to_excel(writer, index=False, sheet_name='Summary')
    return output.getvalue()

# === Streamlit UI ===

st.title("🏠 House Price Predictor (Web App)")

# Data source
source = st.selectbox("Select Data Source", ["Synthetic", "Real", "Upload CSV"])

if source == "Synthetic":
    n = st.slider("Number of Samples", 100, 10000, 1000)
    min_sqft = st.number_input("Min Square Footage", value=500)
    max_sqft = st.number_input("Max Square Footage", value=3500)
    if st.button("Generate & Train Model"):
        df = generate_dataset(n, min_sqft, max_sqft)
elif source == "Real":
    if st.button("Load Real Dataset & Train"):
        df = load_real_dataset()
else:
    uploaded = st.file_uploader("Upload CSV (Bedrooms, SquareFootage, Price)", type=["csv"])
    if uploaded:
        df = pd.read_csv(uploaded)
        if st.button("Train Model"):
            pass
    else:
        df = None

# Train model
if 'df' in locals():
    try:
        model, X_test, y_test, y_pred, mse, r2 = train_model(df)

        st.success(f"✅ Model Trained\n\nMSE: {mse:.2f}, R² Score: {r2:.4f}")

        st.pyplot(plot_results(y_test, y_pred))

        # Export Options
        st.download_button("📥 Download Word Report (.docx)", generate_docx_report(mse, r2), "report.docx")
        st.download_button("📥 Download Excel File (.xlsx)", generate_excel(y_test, y_pred, mse, r2), "results.xlsx")

    except Exception as e:
        st.error(f"Error: {e}")
